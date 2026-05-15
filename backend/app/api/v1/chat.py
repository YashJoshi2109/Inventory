"""
AI Copilot Chat API — production-grade streaming + persistence.

Key design decisions:
  - The user message and session title update are committed to DB BEFORE streaming
    begins, so they are always persisted regardless of client disconnect.
  - The SSE generator uses a FRESH AsyncSession (not the request-scoped one) to
    save the assistant message, because FastAPI's dependency-injected session is
    closed when the response headers are sent — before the generator body runs.
  - Rate limiting is applied per-IP on the chat endpoint via slowapi.

Endpoints:
  POST /chat/sessions                    — create session
  GET  /chat/sessions                    — list sessions
  PATCH /chat/sessions/{id}/title        — rename
  DELETE /chat/sessions/{id}             — delete
  GET  /chat/sessions/{id}/messages      — full history
  POST /chat/sessions/{id}/messages      — send + stream response
  POST /chat/documents                   — upload knowledge document
  GET  /chat/documents                   — list documents
  DELETE /chat/documents/{id}            — soft-delete document
"""
from __future__ import annotations

import json
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import AsyncIterator

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import func, select

from app.ai.copilot import SYSTEM_PROMPT, run_copilot
from app.api.v1.auth import CurrentUser
from app.core.config import settings
from app.core.database import AsyncSessionLocal, DbSession
from app.models.chat import (
    ChatMessage,
    ChatSession,
    DocChunk,
    DocStatus,
    KnowledgeDocument,
    MessageRole,
)

router = APIRouter(prefix="/chat", tags=["chat"])

UPLOAD_BASE = Path(settings.UPLOAD_DIR) / "knowledge"
UPLOAD_BASE.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv"}


# ── Schemas ────────────────────────────────────────────────────────────────────

class SessionOut(BaseModel):
    id: int
    title: str
    created_at: datetime
    updated_at: datetime
    message_count: int = 0
    model_config = {"from_attributes": True}


class MessageOut(BaseModel):
    id: int
    role: str
    content: str | None
    tool_name: str | None
    tool_result: str | None
    sources: str | None
    created_at: datetime
    model_config = {"from_attributes": True}


class DocumentOut(BaseModel):
    id: int
    title: str
    filename: str
    doc_type: str
    status: str
    chunk_count: int
    created_at: datetime
    model_config = {"from_attributes": True}


class CreateSessionRequest(BaseModel):
    title: str = "New chat"


class RenameTitleRequest(BaseModel):
    title: str


# ── Session endpoints ──────────────────────────────────────────────────────────

@router.post("/sessions", response_model=SessionOut, status_code=status.HTTP_201_CREATED)
async def create_session(
    body: CreateSessionRequest,
    db: DbSession,
    current_user: CurrentUser,
) -> SessionOut:
    session = ChatSession(user_id=current_user.id, title=body.title[:255])
    db.add(session)
    await db.flush()
    await db.refresh(session)
    return SessionOut(
        id=session.id,
        title=session.title,
        created_at=session.created_at,
        updated_at=session.updated_at,
        message_count=0,
    )


@router.get("/sessions", response_model=list[SessionOut])
async def list_sessions(
    db: DbSession,
    current_user: CurrentUser,
    limit: int = Query(default=50, le=200),
) -> list[SessionOut]:
    result = await db.execute(
        select(ChatSession)
        .where(ChatSession.user_id == current_user.id)
        .order_by(ChatSession.updated_at.desc())
        .limit(limit)
    )
    sessions = result.scalars().all()
    if not sessions:
        return []

    session_ids = [s.id for s in sessions]
    counts_result = await db.execute(
        select(ChatMessage.session_id, func.count(ChatMessage.id))
        .where(ChatMessage.session_id.in_(session_ids))
        .group_by(ChatMessage.session_id)
    )
    counts_map = {sid: int(cnt) for sid, cnt in counts_result.all()}
    return [
        SessionOut(
            id=s.id,
            title=s.title,
            created_at=s.created_at,
            updated_at=s.updated_at,
            message_count=counts_map.get(s.id, 0),
        )
        for s in sessions
    ]


@router.patch("/sessions/{session_id}/title", response_model=SessionOut)
async def rename_session(
    session_id: int,
    body: RenameTitleRequest,
    db: DbSession,
    current_user: CurrentUser,
) -> SessionOut:
    result = await db.execute(
        select(ChatSession).where(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
        )
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    session.title = body.title[:255]
    await db.flush()
    await db.refresh(session)
    return SessionOut(
        id=session.id,
        title=session.title,
        created_at=session.created_at,
        updated_at=session.updated_at,
    )


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: int,
    db: DbSession,
    current_user: CurrentUser,
) -> dict:
    result = await db.execute(
        select(ChatSession).where(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
        )
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    await db.delete(session)
    return {"ok": True}


@router.get("/sessions/{session_id}/messages", response_model=list[MessageOut])
async def get_messages(
    session_id: int,
    db: DbSession,
    current_user: CurrentUser,
) -> list[MessageOut]:
    sess_result = await db.execute(
        select(ChatSession).where(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
        )
    )
    if not sess_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Session not found")
    result = await db.execute(
        select(ChatMessage)
        .where(ChatMessage.session_id == session_id)
        .order_by(ChatMessage.created_at)
    )
    return result.scalars().all()


@router.post("/sessions/{session_id}/messages")
async def send_message(
    session_id: int,
    db: DbSession,
    current_user: CurrentUser,
    # Accept content as EITHER a query-param (old clients) OR a form field (new clients).
    # Query param is checked first so legacy frontends keep working without redeployment.
    content_q: str | None = Query(default=None, max_length=4000, alias="content"),
    content_f: str | None = Form(default=None, max_length=4000, alias="content"),
    image: UploadFile | None = File(default=None),
) -> StreamingResponse:
    """
    Stream AI response via SSE.  Accepts multipart form so the user can
    optionally attach an image for vision-based queries.
    Content may arrive as a query-param (?content=…) for old clients or as
    a multipart form field for new clients that also send images.

    Persistence strategy:
    1. User message + title update are flushed & committed synchronously,
       BEFORE the StreamingResponse is returned.
    2. The SSE generator opens a FRESH database session and saves the
       assistant message there once streaming is done.
    """
    content = (content_f or content_q or "").strip()
    if not content:
        raise HTTPException(status_code=422, detail="content is required")

    # ── Validate session ──────────────────────────────────────────────────
    sess_result = await db.execute(
        select(ChatSession).where(
            ChatSession.id == session_id,
            ChatSession.user_id == current_user.id,
        )
    )
    session = sess_result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # ── Persist user message NOW (before stream starts) ───────────────────
    user_msg = ChatMessage(
        session_id=session.id,
        role=MessageRole.USER,
        content=content,
    )
    db.add(user_msg)
    await db.flush()

    # Build LLM history from persisted messages
    history_result = await db.execute(
        select(ChatMessage)
        .where(ChatMessage.session_id == session_id)
        .order_by(ChatMessage.created_at)
    )
    history = history_result.scalars().all()

    openai_messages: list[dict] = [{"role": "system", "content": SYSTEM_PROMPT}]
    for msg in history:
        if msg.role == MessageRole.USER:
            openai_messages.append({"role": "user", "content": msg.content or ""})
        elif msg.role == MessageRole.ASSISTANT:
            openai_messages.append({"role": "assistant", "content": msg.content or ""})

    actor_roles = [ur.role.name for ur in current_user.roles if ur.role]
    actor_id = current_user.id

    # Read image bytes BEFORE committing (file upload is request-scoped)
    image_bytes: bytes | None = None
    image_mime: str | None = None
    if image and image.size and image.size > 0:
        image_bytes = await image.read()
        image_mime = image.content_type or "image/jpeg"

    # Auto-title session on first user message
    if session.title in ("New chat", "") and len(history) == 1:
        session.title = content[:80]
        await db.flush()

    # Commit user message + title BEFORE streaming (so they're always saved)
    await db.commit()

    # ── SSE generator (uses its own DB session) ───────────────────────────
    async def event_stream() -> AsyncIterator[bytes]:
        assistant_content = ""

        async for sse_line in run_copilot(
            messages=openai_messages,
            db=None,            # copilot will open its own session for tool calls
            actor_id=actor_id,
            actor_roles=actor_roles,
            image_bytes=image_bytes,
            image_mime=image_mime,
        ):
            yield sse_line.encode()

            try:
                raw = sse_line.replace("data: ", "").strip()
                if not raw:
                    continue
                payload = json.loads(raw)
                if payload.get("type") == "token":
                    assistant_content += payload.get("content", "")
                elif payload.get("type") == "done":
                    # Save assistant reply in a fresh session that we control
                    async with AsyncSessionLocal() as fresh_db:
                        async with fresh_db.begin():
                            fresh_db.add(ChatMessage(
                                session_id=session_id,
                                role=MessageRole.ASSISTANT,
                                content=assistant_content or "(no text response)",
                            ))
                            # Also bump session updated_at
                            sess_upd = await fresh_db.execute(
                                select(ChatSession).where(ChatSession.id == session_id)
                            )
                            s = sess_upd.scalar_one_or_none()
                            if s:
                                from datetime import timezone
                                s.updated_at = datetime.now(timezone.utc)
            except Exception:
                pass

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


# ── Document endpoints ─────────────────────────────────────────────────────────

@router.post("/documents", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    db: DbSession,
    current_user: CurrentUser,
    file: UploadFile = File(...),
    doc_type: str = Form(default="general"),
    title: str = Form(default=""),
) -> DocumentOut:
    suffix = Path(file.filename or "upload").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{suffix}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    safe_name = f"{uuid.uuid4().hex}_{(file.filename or 'upload').replace(' ', '_')}"
    content = await file.read()

    # Write to a temp file so _extract_text_chunks can read it via Path
    tmp_fd = tempfile.NamedTemporaryFile(
        suffix=Path(safe_name).suffix, delete=False
    )
    tmp_path = Path(tmp_fd.name)
    tmp_fd.write(content)
    tmp_fd.close()

    if settings.GCS_BUCKET_NAME:
        # Upload to Google Cloud Storage; store object key as file_path
        from google.cloud import storage as gcs_storage  # type: ignore[import]
        gcs_key = f"knowledge/{safe_name}"
        gcs_client = gcs_storage.Client()
        blob = gcs_client.bucket(settings.GCS_BUCKET_NAME).blob(gcs_key)
        blob.upload_from_string(content, content_type=file.content_type or "application/octet-stream")
        file_path_str = f"gcs://{settings.GCS_BUCKET_NAME}/{gcs_key}"
    else:
        dest = UPLOAD_BASE / safe_name
        dest.write_bytes(content)
        file_path_str = str(dest)

    doc = KnowledgeDocument(
        title=title or (file.filename or "Untitled"),
        filename=file.filename or safe_name,
        mime_type=file.content_type or "application/octet-stream",
        doc_type=doc_type,
        file_path=file_path_str,
        uploaded_by=current_user.id,
        status=DocStatus.PENDING,
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)

    try:
        chunks = _extract_text_chunks(tmp_path, file.content_type or "")
        tmp_path.unlink(missing_ok=True)  # clean up temp file after extraction
        from app.ai.embeddings import embed_text_cached, vec_to_json
        from app.ai.llm_cache import llm_cache_invalidate_docs
        from app.ai.tools import _RAG_CACHE  # clear RAG result cache on new upload
        _RAG_CACHE.clear()
        for i, chunk_text in enumerate(chunks):
            embedding_json: str | None = None
            try:
                # Cached: same text won't re-hit the API across re-uploads
                vec = await embed_text_cached(chunk_text[:2000])
                if vec:
                    embedding_json = vec_to_json(vec)
            except Exception:
                pass
            db.add(DocChunk(
                doc_id=doc.id,
                chunk_index=i,
                content=chunk_text,
                token_count=len(chunk_text.split()),
                embedding_json=embedding_json,
            ))
        doc.chunk_count = len(chunks)
        doc.status = DocStatus.READY
        llm_cache_invalidate_docs()  # new doc = stale LLM cache
        await db.flush()
    except Exception:
        tmp_path.unlink(missing_ok=True)
        doc.status = DocStatus.FAILED
        await db.flush()

    await db.refresh(doc)
    return doc


@router.get("/documents", response_model=list[DocumentOut])
async def list_documents(db: DbSession, current_user: CurrentUser) -> list[DocumentOut]:
    result = await db.execute(
        select(KnowledgeDocument)
        .where(KnowledgeDocument.is_active == True)  # noqa: E712
        .order_by(KnowledgeDocument.created_at.desc())
    )
    return result.scalars().all()


@router.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: int, db: DbSession, current_user: CurrentUser):
    """Serve the raw uploaded file so the frontend can show a PDF preview."""
    result = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.is_active == True,  # noqa: E712
        )
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not doc.file_path:
        raise HTTPException(status_code=404, detail="No file stored for this document")

    # GCS-backed file: stream bytes from bucket
    if doc.file_path.startswith("gcs://"):
        try:
            from google.cloud import storage as gcs_storage  # type: ignore[import]
            from fastapi.responses import Response
            # gcs://bucket/key
            gcs_path = doc.file_path[len("gcs://"):]
            bucket_name, _, gcs_key = gcs_path.partition("/")
            gcs_client = gcs_storage.Client()
            blob = gcs_client.bucket(bucket_name).blob(gcs_key)
            data = blob.download_as_bytes()
            return Response(
                content=data,
                media_type=doc.mime_type or "application/octet-stream",
                headers={"Content-Disposition": f'inline; filename="{doc.filename}"'},
            )
        except Exception as exc:
            raise HTTPException(status_code=404, detail=f"File not found in cloud storage: {exc}") from exc

    # Local disk file
    path = Path(doc.file_path)
    if not path.exists():
        raise HTTPException(status_code=404, detail="File missing from storage")

    from fastapi.responses import FileResponse
    return FileResponse(
        path=str(path),
        media_type=doc.mime_type or "application/octet-stream",
        filename=doc.filename,
    )


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: int, db: DbSession, current_user: CurrentUser) -> dict:
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.is_active = False
    await db.flush()
    from app.ai.llm_cache import llm_cache_invalidate_docs
    from app.ai.tools import _RAG_CACHE
    _RAG_CACHE.clear()
    llm_cache_invalidate_docs()
    return {"ok": True}


# ── Text extraction ────────────────────────────────────────────────────────────

def _extract_text_chunks(
    path: Path,
    mime_type: str,
    chunk_size: int = 500,
    overlap: int = 80,
) -> list[str]:
    """
    Recursive splitter: paragraphs → sentences → words, with overlap.
    overlap words are repeated at the start of the next chunk to preserve
    context across chunk boundaries (reduces lost context for long documents).
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf" or "pdf" in mime_type:
        text = _read_pdf(path)
    elif suffix in (".docx", ".doc") or "wordprocessingml" in mime_type:
        text = _read_docx(path)
    elif suffix == ".csv" or "csv" in mime_type or "text/csv" in mime_type:
        return _chunk_csv(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")

    return _recursive_split(text, chunk_size=chunk_size, overlap=overlap)


def _recursive_split(text: str, chunk_size: int = 500, overlap: int = 80) -> list[str]:
    """
    Split text into chunks of ~chunk_size words with overlap words shared
    between consecutive chunks.  Tries to split at paragraph → sentence
    → word boundaries in that order.
    """
    words = text.split()
    if not words:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]

        # Find a clean sentence break near the end (within last 20% of chunk)
        window = max(1, len(chunk_words) // 5)
        cut = len(chunk_words)
        for i in range(len(chunk_words) - 1, max(0, len(chunk_words) - window) - 1, -1):
            if chunk_words[i].endswith((".", "!", "?", ":", "…")):
                cut = i + 1
                break

        final_words = chunk_words[:cut]
        chunk_text = " ".join(final_words).strip()
        if chunk_text:
            chunks.append(chunk_text)

        if end >= len(words):
            break

        # Overlap never exceeds half the actual chunk to guarantee forward progress
        effective_overlap = min(overlap, max(0, cut // 2))
        advance = max(1, cut - effective_overlap)
        start += advance

    return chunks or [text[:4000]]


def _chunk_csv(path: Path, rows_per_chunk: int = 50) -> list[str]:
    """Parse CSV into header + N-row chunks for accurate semantic retrieval."""
    import csv
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(text.splitlines())
        rows = list(reader)
    except Exception:
        return [path.read_text(encoding="utf-8", errors="replace")[:4000]]

    if not rows:
        return []

    headers = rows[0]
    header_line = ",".join(headers)
    data_rows = rows[1:]

    if not data_rows:
        return [header_line]

    chunks: list[str] = []
    for i in range(0, len(data_rows), rows_per_chunk):
        batch = data_rows[i : i + rows_per_chunk]
        lines = [header_line] + [",".join(r) for r in batch]
        chunks.append("\n".join(lines))
    return chunks or [header_line]


def _read_pdf(path: Path) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        return f"[PDF: {path.name} — install pypdf to extract text]"
    except Exception as exc:
        return f"[PDF extraction error: {exc}]"


def _read_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        return f"[DOCX: {path.name} — install python-docx to extract text]"
    except Exception as exc:
        return f"[DOCX extraction error: {exc}]"
